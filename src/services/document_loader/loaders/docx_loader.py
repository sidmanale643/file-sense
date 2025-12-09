# src/services/document_loader/loaders/docx_loader.py
from docx import Document
from pathlib import Path

class DocxLoader:
    def load_file(self, file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if not path.is_file():
            raise IsADirectoryError(f"Expected a file but got directory: {file_path}")

        try:
            doc = Document(path)
        except Exception as e:
            raise RuntimeError(f"Failed to open DOCX: {e}")

        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # include tables
        for table in doc.tables:
            for row in table.rows:
                text_parts.append("\t".join(cell.text for cell in row.cells))

        return "\n".join(text_parts)
